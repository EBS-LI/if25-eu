# -*- coding: utf-8 -*-
"""
Baut eine angepasste reference.docx fuer Pandoc-DOCX-Export.
Setzt A4, Segoe UI, Projekt-Farben, Fusszeile mit Kontakt + Seitenzahl.
Benoetigt: pip install python-docx
"""
import sys, os, subprocess

# python-docx installieren falls noetig
try:
    from docx import Document
    from docx.shared import Pt, Mm, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Installiere python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Mm, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SRC_DOCX = os.path.join(SCRIPT_DIR, "reference.docx")
OUT_DOCX = os.path.join(SCRIPT_DIR, "reference_biometh.docx")

# Projekt-Farben
PRIMARY   = RGBColor(0x1a, 0x3a, 0x5c)
TEXT_COLOR = RGBColor(0x2c, 0x3e, 0x50)
TEXT_LIGHT = RGBColor(0x5d, 0x6d, 0x7e)
ACCENT    = RGBColor(0x29, 0x80, 0xb9)
FONT_NAME = "Segoe UI"


def set_run_font(run, name=FONT_NAME, size=None, color=None, bold=None):
    """Hilfsfunktion: Font-Eigenschaften auf einen Run setzen."""
    run.font.name = name
    # Auch fuer komplexe Schriften (Asiatisch etc.) setzen
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)
        rFonts.set(qn("w:cs"), name)
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold


def modify_styles(doc):
    """Alle relevanten Styles anpassen."""
    for style in doc.styles:
        # Paragraph- und Character-Styles
        if hasattr(style, "font") and style.font is not None:
            style.font.name = FONT_NAME
            # rFonts fuer alle Varianten
            if hasattr(style, "element"):
                rpr = style.element.find(qn("w:rPr"))
                if rpr is not None:
                    rFonts = rpr.find(qn("w:rFonts"))
                    if rFonts is not None:
                        rFonts.set(qn("w:ascii"), FONT_NAME)
                        rFonts.set(qn("w:hAnsi"), FONT_NAME)
                        rFonts.set(qn("w:cs"), FONT_NAME)

        name = style.name or ""

        # Heading-Styles: Projektfarbe
        if name.startswith("Heading"):
            style.font.color.rgb = PRIMARY
            style.font.name = FONT_NAME
            if name == "Heading 1":
                style.font.size = Pt(16)
            elif name == "Heading 2":
                style.font.size = Pt(13)
            elif name == "Heading 3":
                style.font.size = Pt(11)

        # Normal: 11pt, Textfarbe, Blocksatz
        if name == "Normal":
            style.font.size = Pt(11)
            style.font.color.rgb = TEXT_COLOR
            style.font.name = FONT_NAME
            if style.paragraph_format:
                style.paragraph_format.space_after = Pt(4)
                style.paragraph_format.line_spacing = 1.4
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Body Text / First Paragraph / Compact: auch Blocksatz
        if name in ("Body Text", "First Paragraph", "Compact"):
            if hasattr(style, "paragraph_format") and style.paragraph_format:
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_page_format(doc):
    """A4, Raender."""
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(25)
        section.right_margin = Mm(20)


def add_footer(doc):
    """Einzeilige Fusszeile: Firma | Adresse | Mail ... Tab ... Seite X (rechts)."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Bestehende Absaetze leeren
        for p in footer.paragraphs:
            p.clear()

        # Berechne rechten Tabstopp-Position = Seitenbreite - Raender
        # 210mm - 25mm links - 20mm rechts = 165mm = 9356 Twips (1mm=56.7twips)
        tab_pos_twips = 9356

        # Einzeiliger Footer-Absatz
        p1 = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Obere Rahmenlinie als Trennstrich
        pPr = p1._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="1a3a5c"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # Tabstopp rechts definieren
        tabs = parse_xml(
            f'<w:tabs {nsdecls("w")}>'
            f'  <w:tab w:val="right" w:pos="{tab_pos_twips}"/>'
            f'</w:tabs>'
        )
        pPr.append(tabs)

        # Abstand oben
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="120"/>')
        pPr.append(spacing)

        # Links: Firma | Adresse | Mail
        run1 = p1.add_run("Energieberatung Schwaigkofler GmbH")
        set_run_font(run1, size=Pt(7.5), color=PRIMARY, bold=True)

        run_sep = p1.add_run("  |  Poststrasse 2, FL-9491 Ruggell  |  info@schwaigkofler.org")
        set_run_font(run_sep, size=Pt(7.5), color=TEXT_LIGHT)

        # Tab einfuegen
        run_tab = p1.add_run()
        run_tab._element.append(parse_xml(f'<w:tab {nsdecls("w")}/>'))

        # Rechts: Seite X
        run_s = p1.add_run("Seite ")
        set_run_font(run_s, size=Pt(7.5), color=TEXT_LIGHT)

        # PAGE-Feld
        run_pg1 = p1.add_run()
        run_pg1._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        set_run_font(run_pg1, size=Pt(7.5), color=TEXT_LIGHT)

        run_pg2 = p1.add_run()
        run_pg2._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        set_run_font(run_pg2, size=Pt(7.5), color=TEXT_LIGHT)

        run_pg3 = p1.add_run()
        run_pg3._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
        set_run_font(run_pg3, size=Pt(7.5), color=TEXT_LIGHT)


def main():
    doc = Document(SRC_DOCX)
    modify_styles(doc)
    set_page_format(doc)
    add_footer(doc)
    doc.save(OUT_DOCX)
    print(f"Reference-DOCX erstellt: {OUT_DOCX}")


if __name__ == "__main__":
    main()
